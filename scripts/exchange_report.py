#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
汇率 Excel 报表生成脚本

运行方式：
1. 安装依赖：pip3 install -r requirements.txt
2. 运行脚本：python3 exchange_report.py
3. 或者在 Mac 上双击“生成汇率报表.command”
"""

from __future__ import annotations

import itertools
import subprocess
import sys
import time
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Optional

import requests
from docx import Document
from openpyxl import Workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


PROJECT_DIR = Path(__file__).resolve().parent.parent
REPORT_DIR = PROJECT_DIR / "reports" / "exchange"
TIMEOUT_SECONDS = 15
FRANKFURTER_URL = "https://api.frankfurter.app/{day}"
CURRENCIES = ["USD", "CNY", "MYR", "SGD"]
CURRENCY_NAMES = {
    "USD": "美元",
    "CNY": "人民币",
    "MYR": "马来西亚币",
    "SGD": "新加坡元",
}
PERIOD_DAYS = {
    "realtime": 0,
    "1D": 1,
    "1W": 7,
    "1M": 30,
    "3M": 90,
    "1Y": 365,
    "3Y": 365 * 3,
    "5Y": 365 * 5,
    "10Y": 365 * 10,
}
PERIOD_LABELS = {
    "realtime": "实时",
    "1D": "1日",
    "1W": "1周",
    "1M": "1个月",
    "3M": "3个月",
    "1Y": "1年",
    "3Y": "3年",
    "5Y": "5年",
    "10Y": "10年",
}
SOURCE_NAME = "Frankfurter.app（欧洲央行公开汇率数据）"


@dataclass
class RateSnapshot:
    """保存某一天从接口返回的 USD 基准汇率。"""

    label: str
    requested_date: date
    data_date: str
    rates: dict[str, float]


def build_session() -> requests.Session:
    """创建带浏览器标识的请求会话。"""

    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126 Safari/537.36"
            ),
            "Accept": "application/json,*/*",
        }
    )
    return session


def check_network(session: requests.Session) -> tuple[bool, str]:
    """检查网络是否可用。"""

    try:
        response = session.get("https://www.apple.com/library/test/success.html", timeout=TIMEOUT_SECONDS)
        if response.status_code < 500:
            return True, "网络连接正常"
        return False, f"网络检查返回异常状态码：{response.status_code}"
    except requests.RequestException as exc:
        return False, f"网络检查失败：{exc}"


def request_with_retry(session: requests.Session, url: str, params: dict[str, str], retries: int = 3) -> dict:
    """带简单重试的 JSON 请求，避免临时网络抖动直接导致失败。"""

    last_error: Optional[Exception] = None
    for attempt in range(1, retries + 1):
        try:
            response = session.get(url, params=params, timeout=TIMEOUT_SECONDS)
            response.raise_for_status()
            return response.json()
        except Exception as exc:
            last_error = exc
            print(f"接口请求失败，第 {attempt} 次：{exc}")
            if attempt < retries:
                time.sleep(1.2 * attempt)

    raise RuntimeError(f"接口请求失败，已重试 {retries} 次：{last_error}")


def fetch_usd_snapshot(session: requests.Session, label: str, requested_date: date) -> RateSnapshot:
    """抓取指定日期的 USD 对其他货币汇率。"""

    url = FRANKFURTER_URL.format(day=requested_date.isoformat())
    params = {"from": "USD", "to": "CNY,MYR,SGD"}
    payload = request_with_retry(session, url, params)

    rates = {"USD": 1.0}
    for currency in ["CNY", "MYR", "SGD"]:
        value = (payload.get("rates") or {}).get(currency)
        if not isinstance(value, (int, float)):
            raise RuntimeError(f"接口没有返回 {currency} 汇率：{payload}")
        rates[currency] = float(value)

    data_date = payload.get("date")
    if not data_date:
        raise RuntimeError(f"接口没有返回数据日期：{payload}")

    return RateSnapshot(label=label, requested_date=requested_date, data_date=data_date, rates=rates)


def cross_rate(snapshot: RateSnapshot, base: str, quote: str) -> float:
    """计算货币对汇率，例如 USD/CNY 或 CNY/USD。"""

    return snapshot.rates[quote] / snapshot.rates[base]


def pct_change(current: float, previous: float) -> Optional[float]:
    """计算涨幅。涨幅 = 当前汇率 / 历史汇率 - 1。"""

    if previous == 0:
        return None
    return current / previous - 1


def normalize_currency(value: str) -> str:
    """规范化币种代码。"""

    return (value or "").strip().upper()


def validate_pair(base: str, quote: str) -> tuple[str, str]:
    """检查用户选择的货币对是否合法。"""

    base = normalize_currency(base)
    quote = normalize_currency(quote)
    if base not in CURRENCIES:
        raise ValueError(f"不支持的基准货币：{base}")
    if quote not in CURRENCIES:
        raise ValueError(f"不支持的目标货币：{quote}")
    if base == quote:
        raise ValueError("基准货币和目标货币不能相同")
    return base, quote


def validate_period(period: str) -> str:
    """检查时间范围是否合法。"""

    period = (period or "realtime").strip()
    if period not in PERIOD_DAYS:
        raise ValueError(f"不支持的时间范围：{period}")
    return period


def collect_exchange_rate_pair(base: str, quote: str, period: str = "realtime") -> dict:
    """只查询单个货币对，不生成任何文档。"""

    base, quote = validate_pair(base, quote)
    period = validate_period(period)

    session = build_session()
    network_ok, network_message = check_network(session)
    print(network_message)
    if not network_ok:
        print("网络预检查失败，继续尝试访问汇率接口...")

    today = datetime.now().date()
    current = fetch_usd_snapshot(session, "当前", today)
    current_rate = cross_rate(current, base, quote)

    start_snapshot = None
    start_rate = None
    change_amount = None
    change_percent = None
    requested_start_date = None

    if period != "realtime":
        requested_start_date = today - timedelta(days=PERIOD_DAYS[period])
        start_snapshot = fetch_usd_snapshot(session, PERIOD_LABELS[period], requested_start_date)
        start_rate = cross_rate(start_snapshot, base, quote)
        change_amount = current_rate - start_rate
        change_percent = pct_change(current_rate, start_rate)

    return {
        "base": base,
        "quote": quote,
        "pair": f"{base}/{quote}",
        "period": period,
        "period_label": PERIOD_LABELS[period],
        "current_rate": current_rate,
        "start_rate": start_rate,
        "change_amount": change_amount,
        "change_percent": change_percent,
        "requested_current_date": today.isoformat(),
        "current_date": current.data_date,
        "requested_start_date": requested_start_date.isoformat() if requested_start_date else None,
        "start_date": start_snapshot.data_date if start_snapshot else None,
        "source": SOURCE_NAME,
        "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def collect_exchange_data() -> tuple[list[dict], list[RateSnapshot], str]:
    """抓取汇率数据，并整理成报表行。"""

    session = build_session()
    network_ok, network_message = check_network(session)
    print(network_message)
    if not network_ok:
        print("网络预检查失败，继续尝试访问汇率接口...")

    today = datetime.now().date()
    snapshot_specs = [
        ("当前", today),
        ("前一日", today - timedelta(days=1)),
        ("前一周", today - timedelta(days=7)),
        ("前一月", today - timedelta(days=30)),
    ]

    snapshots = []
    for label, requested_date in snapshot_specs:
        print(f"正在抓取{label}汇率：{requested_date.isoformat()}")
        snapshots.append(fetch_usd_snapshot(session, label, requested_date))

    current, day_before, week_before, month_before = snapshots
    rows = []

    for base, quote in itertools.permutations(CURRENCIES, 2):
        current_rate = cross_rate(current, base, quote)
        day_rate = cross_rate(day_before, base, quote)
        week_rate = cross_rate(week_before, base, quote)
        month_rate = cross_rate(month_before, base, quote)

        rows.append(
            {
                "pair": f"{base}/{quote}",
                "base": base,
                "quote": quote,
                "current_rate": current_rate,
                "day_change": pct_change(current_rate, day_rate),
                "week_change": pct_change(current_rate, week_rate),
                "month_change": pct_change(current_rate, month_rate),
                "current_date": current.data_date,
                "day_date": day_before.data_date,
                "week_date": week_before.data_date,
                "month_date": month_before.data_date,
            }
        )

    return rows, snapshots, SOURCE_NAME


def pair_report_filename(data: dict, suffix: str) -> str:
    """生成单货币对报表文件名。"""

    now = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return f"汇率报表_{data['base']}_{data['quote']}_{data['period']}_{now}{suffix}"


def build_pair_excel_report(data: dict, output_dir: Optional[Path] = None) -> Path:
    """生成单个货币对 Excel 报表。"""

    target_dir = Path(output_dir) if output_dir else REPORT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / pair_report_filename(data, ".xlsx")

    wb = Workbook()
    ws = wb.active
    ws.title = "汇率查询"
    note_ws = wb.create_sheet("说明")

    ws.merge_cells("A1:B1")
    ws["A1"] = "汇率查询报表"
    style_title_cell(ws["A1"])
    ws.row_dimensions[1].height = 28

    rows = [
        ("基准货币", f"{data['base']} {CURRENCY_NAMES[data['base']]}"),
        ("目标货币", f"{data['quote']} {CURRENCY_NAMES[data['quote']]}"),
        ("时间范围", data["period_label"]),
        ("当前汇率", data["current_rate"]),
        ("起始汇率", data["start_rate"] if data["start_rate"] is not None else "实时查询不适用"),
        ("涨跌金额", data["change_amount"] if data["change_amount"] is not None else "实时查询不适用"),
        ("涨跌百分比", data["change_percent"] if data["change_percent"] is not None else "实时查询不适用"),
        ("当前数据日期", data["current_date"]),
        ("起始数据日期", data["start_date"] or "实时查询不适用"),
        ("数据来源", data["source"]),
        ("数据更新时间", data["updated_at"]),
    ]

    ws["A3"] = "指标"
    ws["B3"] = "数值"
    style_header_row(ws, 3, 1, 2)

    for row_index, (label, value) in enumerate(rows, start=4):
        ws.cell(row=row_index, column=1, value=label)
        ws.cell(row=row_index, column=2, value=value)

    for row_index in range(4, 15):
        ws.cell(row=row_index, column=1).font = Font(name="Arial", bold=True)
    for row_index in [7, 8, 9]:
        if isinstance(ws.cell(row=row_index, column=2).value, (int, float)):
            ws.cell(row=row_index, column=2).number_format = "0.000000"
    if isinstance(ws["B10"].value, (int, float)):
        ws["B10"].number_format = "0.00%"
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 42

    note_ws["A1"] = "说明"
    style_title_cell(note_ws["A1"])
    notes = [
        ("当前汇率", f"表示 1 {data['base']} 可兑换多少 {data['quote']}。"),
        ("起始汇率", "当时间范围不是实时查询时，使用所选时间范围起点附近的公开汇率数据。"),
        ("涨跌金额", "涨跌金额 = 当前汇率 - 起始汇率。"),
        ("涨跌百分比", "涨跌百分比 = 当前汇率 / 起始汇率 - 1。"),
        ("数据来源", data["source"]),
    ]
    for row_index, (label, text) in enumerate(notes, start=3):
        note_ws.cell(row=row_index, column=1, value=label)
        note_ws.cell(row=row_index, column=2, value=text)
        note_ws.cell(row=row_index, column=1).font = Font(name="Arial", bold=True)
    note_ws.column_dimensions["A"].width = 16
    note_ws.column_dimensions["B"].width = 70

    for sheet in [ws, note_ws]:
        set_common_sheet_style(sheet)

    wb.save(output_path)
    return output_path


def build_pair_word_report(data: dict, output_dir: Optional[Path] = None) -> Path:
    """生成单个货币对 Word 报表。"""

    target_dir = Path(output_dir) if output_dir else REPORT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / pair_report_filename(data, ".docx")

    document = Document()
    title = document.add_heading("汇率查询报表", level=0)
    title.alignment = 1

    document.add_paragraph(f"生成时间：{data['updated_at']}")
    document.add_paragraph(f"数据来源：{data['source']}")

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "数值"

    def add_row(label: str, value: object) -> None:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = "" if value is None else str(value)

    add_row("基准货币", f"{data['base']} {CURRENCY_NAMES[data['base']]}")
    add_row("目标货币", f"{data['quote']} {CURRENCY_NAMES[data['quote']]}")
    add_row("时间范围", data["period_label"])
    add_row("当前汇率", f"{data['current_rate']:.6f}")
    add_row("起始汇率", f"{data['start_rate']:.6f}" if data["start_rate"] is not None else "实时查询不适用")
    add_row("涨跌金额", f"{data['change_amount']:.6f}" if data["change_amount"] is not None else "实时查询不适用")
    add_row("涨跌百分比", f"{data['change_percent']:.2%}" if data["change_percent"] is not None else "实时查询不适用")
    add_row("当前数据日期", data["current_date"])
    add_row("起始数据日期", data["start_date"] or "实时查询不适用")
    add_row("数据更新时间", data["updated_at"])

    document.add_heading("说明", level=1)
    document.add_paragraph(f"当前汇率表示 1 {data['base']} 可兑换多少 {data['quote']}。")
    if data["period"] != "realtime":
        document.add_paragraph("涨跌金额 = 当前汇率 - 起始汇率。")
        document.add_paragraph("涨跌百分比 = 当前汇率 / 起始汇率 - 1。")
    document.add_paragraph("本报表仅供参考，不构成投资建议。")

    document.save(output_path)
    return output_path


def style_title_cell(cell) -> None:
    """设置标题单元格样式。"""

    cell.font = Font(name="Arial", bold=True, size=18, color="1F4E79")
    cell.alignment = Alignment(horizontal="center", vertical="center")


def style_header_row(ws, row_number: int, start_col: int, end_col: int) -> None:
    """设置表头样式。"""

    fill = PatternFill("solid", fgColor="D9EAF7")
    border = Border(bottom=Side(style="thin", color="808080"))
    for col in range(start_col, end_col + 1):
        cell = ws.cell(row=row_number, column=col)
        cell.font = Font(name="Arial", bold=True, color="000000")
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center")


def set_common_sheet_style(ws) -> None:
    """设置工作表通用显示效果。"""

    ws.sheet_view.showGridLines = False
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            old_font = cell.font
            color = None
            if old_font and old_font.color and old_font.color.type == "rgb":
                color = old_font.color.rgb
            cell.font = Font(
                name="Arial",
                size=old_font.sz if old_font and old_font.sz else 11,
                bold=bool(old_font and old_font.bold),
                italic=bool(old_font and old_font.italic),
                color=color,
            )
            cell.alignment = Alignment(
                horizontal=cell.alignment.horizontal,
                vertical="center",
                wrap_text=cell.alignment.wrap_text,
            )
            cell.border = Border(bottom=Side(style="hair", color="D9D9D9"))


def add_change_color_rules(ws, cell_range: str) -> None:
    """给涨跌幅区域增加条件格式：上涨绿色，下跌红色。"""

    green_fill = PatternFill("solid", fgColor="E2F0D9")
    red_fill = PatternFill("solid", fgColor="FCE4D6")
    green_font = Font(color="006100")
    red_font = Font(color="9C0006")
    ws.conditional_formatting.add(
        cell_range,
        CellIsRule(operator="greaterThan", formula=["0"], fill=green_fill, font=green_font),
    )
    ws.conditional_formatting.add(
        cell_range,
        CellIsRule(operator="lessThan", formula=["0"], fill=red_fill, font=red_font),
    )


def build_workbook(rows: list[dict], snapshots: list[RateSnapshot], source: str, output_dir: Optional[Path] = None) -> Path:
    """生成 Excel 报表。"""

    now = datetime.now()
    target_dir = Path(output_dir) if output_dir else REPORT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"汇率报表_{now.strftime('%Y-%m-%d_%H-%M')}.xlsx"

    wb = Workbook()
    ws = wb.active
    ws.title = "汇率报表"
    raw_ws = wb.create_sheet("原始数据")
    note_ws = wb.create_sheet("说明")

    # 主报表
    ws.merge_cells("A1:G1")
    ws["A1"] = "汇率 Excel 报表"
    style_title_cell(ws["A1"])
    ws.row_dimensions[1].height = 28

    ws["A2"] = "生成时间"
    ws["B2"] = now.strftime("%Y-%m-%d %H:%M:%S")
    ws["A3"] = "数据来源"
    ws["B3"] = source
    ws["A4"] = "当前数据日期"
    ws["B4"] = snapshots[0].data_date
    ws["A5"] = "币种范围"
    ws["B5"] = "、".join(f"{code} {CURRENCY_NAMES[code]}" for code in CURRENCIES)

    headers = ["货币对", "基础货币", "报价货币", "当前汇率", "日涨幅", "周涨幅", "月涨幅"]
    for col, header in enumerate(headers, start=1):
        ws.cell(row=7, column=col, value=header)
    style_header_row(ws, 7, 1, len(headers))

    for row_index, row_data in enumerate(rows, start=8):
        ws.cell(row=row_index, column=1, value=row_data["pair"])
        ws.cell(row=row_index, column=2, value=f"{row_data['base']} {CURRENCY_NAMES[row_data['base']]}")
        ws.cell(row=row_index, column=3, value=f"{row_data['quote']} {CURRENCY_NAMES[row_data['quote']]}")
        ws.cell(row=row_index, column=4, value=row_data["current_rate"])
        ws.cell(row=row_index, column=5, value=row_data["day_change"])
        ws.cell(row=row_index, column=6, value=row_data["week_change"])
        ws.cell(row=row_index, column=7, value=row_data["month_change"])

    last_row = 7 + len(rows)
    for row in range(8, last_row + 1):
        ws.cell(row=row, column=4).number_format = "0.000000"
        ws.cell(row=row, column=4).alignment = Alignment(horizontal="right", vertical="center")
        for col in [5, 6, 7]:
            ws.cell(row=row, column=col).number_format = "0.00%"
            ws.cell(row=row, column=col).alignment = Alignment(horizontal="right", vertical="center")

    add_change_color_rules(ws, f"E8:G{last_row}")
    ws.freeze_panes = "A8"
    ws.auto_filter.ref = f"A7:G{last_row}"
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 12
    ws.column_dimensions["G"].width = 12

    # 原始数据表
    raw_headers = ["口径", "请求日期", "实际数据日期", "USD", "CNY", "MYR", "SGD"]
    for col, header in enumerate(raw_headers, start=1):
        raw_ws.cell(row=1, column=col, value=header)
    style_header_row(raw_ws, 1, 1, len(raw_headers))

    for row_index, snapshot in enumerate(snapshots, start=2):
        raw_ws.cell(row=row_index, column=1, value=snapshot.label)
        raw_ws.cell(row=row_index, column=2, value=snapshot.requested_date.isoformat())
        raw_ws.cell(row=row_index, column=3, value=snapshot.data_date)
        for col, currency in enumerate(CURRENCIES, start=4):
            raw_ws.cell(row=row_index, column=col, value=snapshot.rates[currency])
            raw_ws.cell(row=row_index, column=col).number_format = "0.000000"

    raw_ws["A7"] = "说明"
    raw_ws["A8"] = "本表以 USD 为统一基准。任意货币对 A/B 的汇率 = USD/B 汇率 / USD/A 汇率。"
    raw_ws["A9"] = "来源 URL"
    raw_ws["B9"] = "https://api.frankfurter.app/"
    raw_ws.column_dimensions["A"].width = 16
    raw_ws.column_dimensions["B"].width = 16
    raw_ws.column_dimensions["C"].width = 16
    for col in ["D", "E", "F", "G"]:
        raw_ws.column_dimensions[col].width = 14

    # 说明表
    note_ws.merge_cells("A1:D1")
    note_ws["A1"] = "说明"
    style_title_cell(note_ws["A1"])
    notes = [
        ("币种", "本报表包含 USD 美元、CNY 人民币、MYR 马来西亚币、SGD 新加坡元。"),
        ("货币对", "例如 USD/CNY 表示 1 USD 可以兑换多少 CNY。"),
        ("当前汇率", "使用接口返回的当前数据日期对应汇率。周末或节假日时，接口可能返回最近一个交易日。"),
        ("日涨幅", "日涨幅 = 当前汇率 / 前一日汇率 - 1。"),
        ("周涨幅", "周涨幅 = 当前汇率 / 前一周汇率 - 1。"),
        ("月涨幅", "月涨幅 = 当前汇率 / 前一月汇率 - 1。"),
        ("颜色", "涨幅大于 0 显示绿色，涨幅小于 0 显示红色。"),
        ("数据来源", source + "，接口地址：https://api.frankfurter.app/"),
    ]
    for row_index, (label, text) in enumerate(notes, start=3):
        note_ws.cell(row=row_index, column=1, value=label)
        note_ws.cell(row=row_index, column=2, value=text)
        note_ws.cell(row=row_index, column=1).font = Font(name="Arial", bold=True)
        note_ws.cell(row=row_index, column=2).alignment = Alignment(wrap_text=True, vertical="center")
    note_ws.column_dimensions["A"].width = 14
    note_ws.column_dimensions["B"].width = 95

    for sheet in [ws, raw_ws, note_ws]:
        set_common_sheet_style(sheet)

    wb.save(output_path)
    return output_path


def open_file_on_mac(path: Path) -> None:
    """在 macOS 上自动打开生成的 Excel 文件。"""

    if sys.platform == "darwin":
        try:
            subprocess.Popen(["open", str(path)])
            print("已尝试自动打开 Excel 文件。")
        except Exception as exc:
            print(f"自动打开失败，可以手动打开文件：{exc}")


def main() -> int:
    """主流程：抓取汇率、生成 Excel、自动打开。"""

    try:
        print("开始抓取汇率数据...")
        rows, snapshots, source = collect_exchange_data()
        print("开始生成 Excel 报表...")
        output_path = build_workbook(rows, snapshots, source)
        print("完成。Excel 文件路径：")
        print(output_path)
    except Exception as exc:
        print()
        print("Excel 报表生成失败。")
        print(f"失败原因：{exc}")
        print("请检查网络连接、汇率接口是否可访问，或稍后再试。")
        return 1

    open_file_on_mac(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
