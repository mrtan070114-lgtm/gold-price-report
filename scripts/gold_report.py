#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
金价抓取 + Word 报表生成脚本

运行方式：
1. 安装依赖：pip3 install -r requirements.txt
2. 运行脚本：python3 gold_report.py
3. 或者在 Mac 上双击“生成金价报表.command”
"""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from datetime import datetime
from io import StringIO
from pathlib import Path
from typing import Optional
from zoneinfo import ZoneInfo

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent.parent
REPORT_DIR = PROJECT_DIR / "reports" / "gold"
YAHOO_HOSTS = ["query1.finance.yahoo.com", "query2.finance.yahoo.com"]
YAHOO_CHART_URL = "https://{host}/v8/finance/chart/{symbol}"
STOOQ_CSV_URL = "https://stooq.com/q/l/"
SWISSQUOTE_XAU_URL = "https://forex-data-feed.swissquote.com/public-quotes/bboquotes/instrument/XAU/USD"
EXCHANGE_RATE_URL = "https://open.er-api.com/v6/latest/USD"
TIMEOUT_SECONDS = 15
TROY_OUNCE_GRAMS = 31.1035


@dataclass
class GoldData:
    """保存报表需要用到的行情数据。"""

    success: bool
    source: str
    update_time: str
    current_price: Optional[float] = None
    change_amount: Optional[float] = None
    change_percent: Optional[float] = None
    open_price: Optional[float] = None
    high_price: Optional[float] = None
    low_price: Optional[float] = None
    previous_close: Optional[float] = None
    usd_cny: Optional[float] = None
    cny_per_gram: Optional[float] = None
    error_message: str = ""


def build_session() -> requests.Session:
    """创建带浏览器标识的请求会话，减少被接口拒绝的概率。"""

    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126 Safari/537.36"
            ),
            "Accept": "application/json,text/csv,text/plain,*/*",
        }
    )
    return session


def check_network(session: requests.Session) -> tuple[bool, str]:
    """检查基础网络是否可用。检查失败不代表一定不能抓取，只作为提示。"""

    test_url = "https://www.apple.com/library/test/success.html"
    try:
        response = session.get(test_url, timeout=TIMEOUT_SECONDS)
        if response.status_code < 500:
            return True, "网络连接正常"
        return False, f"网络检查返回异常状态码：{response.status_code}"
    except requests.RequestException as exc:
        return False, f"网络检查失败：{exc}"


def first_number(values: list[Optional[float]]) -> Optional[float]:
    """从列表中取第一个有效数字。"""

    for value in values:
        if isinstance(value, (int, float)):
            return float(value)
    return None


def last_number(values: list[Optional[float]]) -> Optional[float]:
    """从列表中取最后一个有效数字。"""

    for value in reversed(values):
        if isinstance(value, (int, float)):
            return float(value)
    return None


def format_timestamp(timestamp: Optional[int], timezone_name: str = "") -> str:
    """把接口时间戳转换成可读时间。"""

    if not timestamp:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    try:
        zone = ZoneInfo(timezone_name) if timezone_name else None
        return datetime.fromtimestamp(timestamp, tz=zone).strftime("%Y-%m-%d %H:%M:%S %Z")
    except Exception:
        return datetime.fromtimestamp(timestamp).strftime("%Y-%m-%d %H:%M:%S")


def fetch_yahoo_chart(session: requests.Session, symbol: str, range_value: str, interval: str) -> dict:
    """从 Yahoo Finance 抓取单个品种的图表数据。"""

    last_error = ""
    for host in YAHOO_HOSTS:
        try:
            url = YAHOO_CHART_URL.format(host=host, symbol=symbol)
            response = session.get(
                url,
                params={"range": range_value, "interval": interval},
                timeout=TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            payload = response.json()
            break
        except Exception as exc:
            last_error = f"{host}: {exc}"
    else:
        raise RuntimeError(last_error)

    chart = payload.get("chart", {})
    error = chart.get("error")
    if error:
        raise RuntimeError(json.dumps(error, ensure_ascii=False))

    results = chart.get("result") or []
    if not results:
        raise RuntimeError("Yahoo Finance 没有返回行情数据")

    return results[0]


def fetch_yahoo_gold(session: requests.Session) -> GoldData:
    """优先从 Yahoo Finance 抓取黄金期货 GC=F。"""

    # 分钟级数据更“新”，但有时会被 Yahoo 限流；失败后改用最近 5 个交易日的日线数据。
    last_error = ""
    for range_value, interval in [("1d", "1m"), ("5d", "1d")]:
        try:
            result = fetch_yahoo_chart(session, "GC=F", range_value, interval)
            break
        except Exception as exc:
            last_error = f"{range_value}/{interval}: {exc}"
    else:
        raise RuntimeError(last_error)

    meta = result.get("meta", {})
    quote = (result.get("indicators", {}).get("quote") or [{}])[0]

    close_values = quote.get("close") or []
    open_values = quote.get("open") or []
    high_values = quote.get("high") or []
    low_values = quote.get("low") or []

    current_price = meta.get("regularMarketPrice") or last_number(close_values)
    previous_close = meta.get("chartPreviousClose") or meta.get("previousClose")
    open_price = meta.get("regularMarketOpen") or last_number(open_values)
    high_price = meta.get("regularMarketDayHigh") or last_number(high_values)
    low_price = meta.get("regularMarketDayLow") or last_number(low_values)

    change_amount = None
    change_percent = None
    if isinstance(current_price, (int, float)) and isinstance(previous_close, (int, float)) and previous_close:
        change_amount = float(current_price) - float(previous_close)
        change_percent = change_amount / float(previous_close) * 100

    timestamps = result.get("timestamp") or []
    update_time = format_timestamp(
        timestamps[-1] if timestamps else meta.get("regularMarketTime"),
        meta.get("exchangeTimezoneName", ""),
    )

    return GoldData(
        success=True,
        source=f"Yahoo Finance（GC=F 黄金期货，{interval}）",
        update_time=update_time,
        current_price=float(current_price) if isinstance(current_price, (int, float)) else None,
        change_amount=change_amount,
        change_percent=change_percent,
        open_price=float(open_price) if isinstance(open_price, (int, float)) else None,
        high_price=float(high_price) if isinstance(high_price, (int, float)) else None,
        low_price=float(low_price) if isinstance(low_price, (int, float)) else None,
        previous_close=float(previous_close) if isinstance(previous_close, (int, float)) else None,
    )


def fetch_stooq_gold(session: requests.Session) -> GoldData:
    """备用方案：尝试从 Stooq 抓取黄金数据。字段可能比 Yahoo 少。"""

    last_error = ""
    for symbol in ["gc.f", "xauusd"]:
        try:
            response = session.get(
                STOOQ_CSV_URL,
                params={"s": symbol, "f": "sd2t2ohlc", "h": "", "e": "csv"},
                timeout=TIMEOUT_SECONDS,
            )
            response.raise_for_status()
            rows = list(csv.DictReader(StringIO(response.text)))
            if not rows:
                raise RuntimeError("Stooq 没有返回 CSV 数据")

            row = rows[0]
            close_text = row.get("Close", "")
            if not close_text or close_text == "N/D":
                raise RuntimeError(f"Stooq 返回无效价格：{row}")

            current_price = float(close_text)
            open_price = float(row["Open"]) if row.get("Open") not in ("", "N/D", None) else None
            high_price = float(row["High"]) if row.get("High") not in ("", "N/D", None) else None
            low_price = float(row["Low"]) if row.get("Low") not in ("", "N/D", None) else None
            update_time = f"{row.get('Date', '')} {row.get('Time', '')}".strip()

            return GoldData(
                success=True,
                source=f"Stooq（{symbol}，备用数据源）",
                update_time=update_time,
                current_price=current_price,
                open_price=open_price,
                high_price=high_price,
                low_price=low_price,
            )
        except Exception as exc:
            last_error = f"{symbol}: {exc}"

    raise RuntimeError(f"Stooq 备用数据源也失败：{last_error}")


def fetch_swissquote_gold(session: requests.Session) -> GoldData:
    """备用方案：从 Swissquote 抓取 XAU/USD 现货买卖价，并取中间价。"""

    response = session.get(SWISSQUOTE_XAU_URL, timeout=TIMEOUT_SECONDS)
    response.raise_for_status()
    payload = response.json()

    if not isinstance(payload, list) or not payload:
        raise RuntimeError("Swissquote 没有返回有效数据")

    for item in payload:
        prices = item.get("spreadProfilePrices") or []
        if not prices:
            continue

        price = prices[0]
        bid = price.get("bid")
        ask = price.get("ask")
        if isinstance(bid, (int, float)) and isinstance(ask, (int, float)):
            current_price = (float(bid) + float(ask)) / 2
            ts = item.get("ts")
            update_time = (
                datetime.fromtimestamp(ts / 1000).strftime("%Y-%m-%d %H:%M:%S")
                if isinstance(ts, (int, float))
                else datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )
            return GoldData(
                success=True,
                source="Swissquote（XAU/USD 现货，备用数据源）",
                update_time=update_time,
                current_price=current_price,
            )

    raise RuntimeError("Swissquote 没有返回有效买卖价")


def fetch_usd_cny(session: requests.Session) -> tuple[Optional[float], str]:
    """从 Yahoo Finance 抓取 USD/CNY 汇率，用于估算人民币每克价格。"""

    try:
        last_error = ""
        for range_value, interval in [("1d", "5m"), ("5d", "1d")]:
            try:
                result = fetch_yahoo_chart(session, "CNY=X", range_value, interval)
                break
            except Exception as exc:
                last_error = f"{range_value}/{interval}: {exc}"
        else:
            raise RuntimeError(last_error)

        meta = result.get("meta", {})
        quote = (result.get("indicators", {}).get("quote") or [{}])[0]
        close_values = quote.get("close") or []
        price = meta.get("regularMarketPrice") or last_number(close_values)
        if isinstance(price, (int, float)):
            return float(price), ""
        raise RuntimeError("Yahoo Finance 未返回 USD/CNY 有效价格")
    except Exception as exc:
        yahoo_error = f"Yahoo Finance USD/CNY 汇率抓取失败：{exc}"

    try:
        response = session.get(EXCHANGE_RATE_URL, timeout=TIMEOUT_SECONDS)
        response.raise_for_status()
        payload = response.json()
        cny = (payload.get("rates") or {}).get("CNY")
        if isinstance(cny, (int, float)):
            return float(cny), ""
        return None, yahoo_error + "；开放汇率接口未返回 CNY"
    except Exception as exc:
        return None, yahoo_error + f"；开放汇率接口也失败：{exc}"


def fetch_market_data() -> GoldData:
    """抓取行情数据。先用 Yahoo Finance，失败后用 Stooq。"""

    session = build_session()
    network_ok, network_message = check_network(session)
    print(network_message)

    errors = []
    if not network_ok:
        errors.append(network_message)

    try:
        data = fetch_yahoo_gold(session)
        print("黄金数据来源：Yahoo Finance")
    except Exception as exc:
        yahoo_error = f"Yahoo Finance 抓取失败：{exc}"
        print(yahoo_error)
        errors.append(yahoo_error)

        try:
            data = fetch_swissquote_gold(session)
            print("黄金数据来源：Swissquote 备用数据源")
        except Exception as swissquote_exc:
            swissquote_error = f"Swissquote 备用数据源失败：{swissquote_exc}"
            print(swissquote_error)
            errors.append(swissquote_error)

            try:
                data = fetch_stooq_gold(session)
                print("黄金数据来源：Stooq 备用数据源")
            except Exception as fallback_exc:
                fallback_error = str(fallback_exc)
                print(fallback_error)
                errors.append(fallback_error)
                return GoldData(
                    success=False,
                    source="Yahoo Finance / Swissquote / Stooq",
                    update_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    error_message="\n".join(errors),
                )

    usd_cny, currency_error = fetch_usd_cny(session)
    if usd_cny:
        data.usd_cny = usd_cny
        if data.current_price:
            data.cny_per_gram = data.current_price * usd_cny / TROY_OUNCE_GRAMS
    elif currency_error:
        print(currency_error)
        data.error_message = currency_error

    return data


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """设置表格单元格文字，并指定中文字体。"""

    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, size=10.5)


def set_run_font(run, size: float = 11, bold: Optional[bool] = None, color: Optional[RGBColor] = None) -> None:
    """统一设置中英文字体，避免 Word 里中文显示不稳定。"""

    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    """设置表头背景色。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def money(value: Optional[float], digits: int = 2) -> str:
    """格式化普通数字。"""

    if value is None:
        return "--"
    return f"{value:,.{digits}f}"


def percent(value: Optional[float]) -> str:
    """格式化百分比。"""

    if value is None:
        return "--"
    return f"{value:+.2f}%"


def signed_money(value: Optional[float]) -> str:
    """格式化带正负号的涨跌金额。"""

    if value is None:
        return "--"
    return f"{value:+,.2f}"


def make_analysis(data: GoldData) -> str:
    """根据涨跌幅生成简短分析。"""

    if not data.success:
        return "数据抓取失败，请检查网络或接口。本报表仅供参考，不构成投资建议。"

    if data.change_percent is None or abs(data.change_percent) < 0.05:
        trend = "今日黄金价格波动较小。"
    elif data.change_percent > 0:
        trend = "今日黄金价格上涨。"
    else:
        trend = "今日黄金价格下跌。"

    return trend + " 本报表仅供参考，不构成投资建议。"


def add_paragraph(document: Document, text: str, size: float = 11, bold: bool = False) -> None:
    """添加普通段落。"""

    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def build_docx_report(data: GoldData, output_dir: Optional[Path] = None) -> Path:
    """生成 Word 报表，并返回保存路径。"""

    target_dir = Path(output_dir) if output_dir else REPORT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    now = datetime.now()
    output_path = target_dir / f"金价报表_{now.strftime('%Y-%m-%d_%H-%M')}.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("每日金价报表")
    set_run_font(title_run, size=20, bold=True, color=RGBColor(31, 78, 121))

    add_paragraph(document, "基本信息", size=13, bold=True)
    add_paragraph(document, f"生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')}")
    add_paragraph(document, f"数据来源：{data.source}")
    add_paragraph(document, f"数据更新时间：{data.update_time}")

    if not data.success:
        add_paragraph(document, "数据抓取失败，请检查网络或接口。", size=12, bold=True)
        if data.error_message:
            add_paragraph(document, "错误原因：" + data.error_message.replace("\n", "；"))

    add_paragraph(document, "核心行情表格", size=13, bold=True)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    set_cell_text(header_cells[0], "指标", bold=True)
    set_cell_text(header_cells[1], "数值", bold=True)
    set_cell_shading(header_cells[0], "D9EAF7")
    set_cell_shading(header_cells[1], "D9EAF7")

    rows = [
        ("当前黄金价格 USD/oz", money(data.current_price)),
        ("涨跌金额", signed_money(data.change_amount)),
        ("涨跌幅", percent(data.change_percent)),
        ("开盘价", money(data.open_price)),
        ("最高价", money(data.high_price)),
        ("最低价", money(data.low_price)),
        ("昨日收盘价", money(data.previous_close)),
        ("USD/CNY 汇率", money(data.usd_cny, digits=4)),
        ("估算人民币/克", money(data.cny_per_gram)),
    ]

    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], value)

    add_paragraph(document, "简短分析", size=13, bold=True)
    add_paragraph(document, make_analysis(data))

    if data.error_message and data.success:
        add_paragraph(document, "补充提示", size=13, bold=True)
        add_paragraph(document, data.error_message.replace("\n", "；"))

    document.save(output_path)
    return output_path


def main() -> int:
    """主流程：抓取数据、生成报表、打印结果。"""

    print("开始抓取最新金价...")
    data = fetch_market_data()

    print("开始生成 Word 报表...")
    output_path = build_docx_report(data)

    print("完成。报表文件路径：")
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
